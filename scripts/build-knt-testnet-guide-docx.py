from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/KNT_Testnet_User_Admin_Guide.docx"


TOKENS = {
    "font": "Calibri",
    "east_asia_font": "Microsoft YaHei",
    "title": RGBColor(31, 58, 95),
    "h1": RGBColor(46, 116, 181),
    "h2": RGBColor(31, 77, 120),
    "muted": RGBColor(85, 85, 85),
    "table_header": "F2F4F7",
    "callout": "F4F6F9",
    "warning": "FFF2CC",
    "border": "D9E2F3",
}


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = TOKENS["font"]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["east_asia_font"])
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=11, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width(table, dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(dxa))


def fixed_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = Inches(width)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        shade_cell(hdr[i], TOKENS["table_header"])
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=9.5)
    doc.add_paragraph()
    return table


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, size=24, bold=True, color=TOKENS["title"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=11, color=TOKENS["muted"])


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color=TOKENS["h1"])


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=TOKENS["h2"])


def para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11)


def callout(doc, label, text, fill=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color=TOKENS["border"])
    cell = table.cell(0, 0)
    shade_cell(cell, fill or TOKENS["callout"])
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, bold=True, color=TOKENS["title"])
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = TOKENS["font"]
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["east_asia_font"])
    styles["Normal"].font.size = Pt(11)

    add_title(doc, "KNT 测试网使用教程", "普通用户 / Admin / Keeper 快速操作手册 · BSC Testnet · 2026-05-19")
    callout(
        doc,
        "测试网入口",
        "Admin 控制台：https://knights-admin.dappweb.workers.dev。普通用户没有独立 UI，测试主要通过钱包、PancakeSwap 测试网和 Admin 查询完成。",
    )

    add_h1(doc, "1. 测试网基础信息")
    fixed_table(
        doc,
        ["项目", "测试网配置"],
        [
            ["网络", "BSC Testnet"],
            ["Chain ID", "97"],
            ["区块浏览器", "https://testnet.bscscan.com"],
            ["Admin 控制台", "https://knights-admin.dappweb.workers.dev"],
            ["KNT Proxy", "0x6d15543f858E185240841015159427957f390eAF"],
            ["USDT", "0xacD944e910952c020eb129C50921f180c62c3291"],
            ["LABUBU", "0xe8214E6580f81835F13Ac358DCc72ac7d4d78052"],
            ["KNT/LABUBU Pair", "0x8A5E41D891b6BE9E470Ab3307dE8E24C8073ff00"],
            ["LABUBU/USDT Pair", "0x628578cACC1210f8a082bDe659072aCB4Ad11DEA"],
            ["Pancake V2 Router", "0xD99D1c33F9fC3444f8101754aBC46c52416550D1"],
        ],
        [1.55, 4.85],
    )

    add_h1(doc, "2. 普通用户测试流程")
    callout(
        doc,
        "用户原则",
        "普通用户不需要打开 Admin，也不需要手动调用合约函数。用户只做三件事：钱包转账、Pancake 买卖、把交易哈希给运营核对。",
    )

    add_h2(doc, "2.1 钱包准备")
    numbered(doc, "在 MetaMask 或兼容钱包中切换到 BSC Testnet，确认 Chain ID 为 97。")
    numbered(doc, "准备测试 BNB 用于 gas。")
    numbered(doc, "添加 KNT、USDT、LABUBU 三个测试网代币地址。")
    numbered(doc, "所有测试交易都必须确认网络是 BSC Testnet，不要在主网地址上测试。")

    add_h2(doc, "2.2 绑定推荐人")
    callout(
        doc,
        "前提",
        "Admin 必须先把 Referral Signal KNT 设置为大于 0。当前如果该值为 0，互转 KNT 只会成为普通转账，不会触发推荐绑定。",
        TOKENS["warning"],
    )
    numbered(doc, "推荐人向用户转入固定数量 KNT，数量必须等于 Admin 设置的 Referral Signal KNT。")
    numbered(doc, "用户再向推荐人转回完全相同数量 KNT。")
    numbered(doc, "等待交易确认后，由 Admin 在财务页输入用户地址，检查 referrer 是否已绑定。")
    bullet(doc, "推荐关系通常只能绑定一次，已绑定用户不能随意更换上级。")
    bullet(doc, "互转数量必须完全一致；多转、少转或 Referral Signal 为 0 都可能不绑定。")

    add_h2(doc, "2.3 入金")
    numbered(doc, "用户在钱包中选择测试网 USDT。")
    numbered(doc, "把 USDT 转到 KNT Proxy 地址：0x6d15543f858E185240841015159427957f390eAF。")
    numbered(doc, "保存交易哈希。")
    numbered(doc, "等待 Keeper 自动扫描并处理入金。处理完成后，用户会获得对应 LP/算力记录。")
    bullet(doc, "入金不是 USDT 到账后立即完成，取决于确认数、Keeper 扫描间隔和链上拥堵。")
    bullet(doc, "如果长时间未处理，把交易哈希发给 Admin 查日志。")

    add_h2(doc, "2.4 在 Pancake 买卖 KNT")
    numbered(doc, "打开 PancakeSwap 测试网，并确认钱包网络为 BSC Testnet。")
    numbered(doc, "使用 KNT/LABUBU 交易对买卖 KNT。必要时手动导入 KNT 和 LABUBU 地址。")
    numbered(doc, "买入后等待 Keeper/市场维护更新价格、LP 和奖励状态。")
    numbered(doc, "卖出 KNT 会按合约规则计算卖出税、盈利税和砸盘税，测试时以链上结果为准。")
    bullet(doc, "当前测试网池子价格约为 1 KNT = 3 LABUBU，且 LABUBU/USDT 池约为 1 LABUBU = 1 USDT。")

    add_h2(doc, "2.5 获得奖励和释放币")
    numbered(doc, "普通奖励由合约和 Keeper 周期性推进，用户无需主动调用合约。")
    numbered(doc, "迁移锁仓释放可由 Keeper/Admin 批量代领分发，用户不需要自己调用领取函数。")
    numbered(doc, "用户最终在钱包中看到 KNT 到账，或在 Admin 财务页看到奖励/释放记录。")
    bullet(doc, "测试网已导入 1215 条迁移锁仓，锁仓总量为 351622.272577569467629156 KNT。")

    add_h2(doc, "2.6 用户测试验收清单")
    bullet(doc, "钱包网络为 BSC Testnet。")
    bullet(doc, "KNT、USDT、LABUBU 地址正确。")
    bullet(doc, "推荐互转金额等于 Referral Signal KNT，且用户未绑定过上级。")
    bullet(doc, "USDT 入金交易成功，并已超过 Keeper 设置的确认数。")
    bullet(doc, "Pancake 买卖使用 KNT/LABUBU 测试网交易对。")
    bullet(doc, "异常反馈必须带用户地址、交易哈希、发生时间和操作说明。")

    add_h1(doc, "3. Admin / Keeper 测试流程")
    add_h2(doc, "3.1 登录和状态检查")
    numbered(doc, "打开 Admin 控制台：https://knights-admin.dappweb.workers.dev。")
    numbered(doc, "连接有权限的钱包，确认页面显示 bscTestnet、Chain ID 97 和新 KNT Proxy。")
    numbered(doc, "进入总览页，检查价格、LP、rewardPool、nodeCount、migration rows 是否正常。")
    numbered(doc, "当前基准状态：nodeCount = 328，migration rows = 1215，totalPower = 65346806.123652943211440508。")

    add_h2(doc, "3.2 Admin 首次配置检查")
    fixed_table(
        doc,
        ["检查项", "期望值 / 操作"],
        [
            ["KNT Proxy", "0x6d15543f858E185240841015159427957f390eAF"],
            ["USDT", "0xacD944e910952c020eb129C50921f180c62c3291"],
            ["LABUBU", "0xe8214E6580f81835F13Ac358DCc72ac7d4d78052"],
            ["KNT Price", "已执行 maintenance 后应为 3.0 USDT"],
            ["Global LP", "已执行 maintenance 后应为 8000.0 USDT"],
            ["Referral Signal KNT", "如果要测试推荐互转绑定，必须设置为大于 0"],
            ["Keeper 钱包", "必须有 Keeper/Manager/Admin/Owner 角色，并有测试 BNB"],
        ],
        [2.0, 4.4],
    )

    add_h2(doc, "3.3 Keeper 手动测试顺序")
    numbered(doc, "运行 Observer：扫描 USDT 入金事件，不直接发链上交易。")
    numbered(doc, "运行 Deposit：处理已发现且满足确认数的 USDT 入金。")
    numbered(doc, "运行 LP Sync：同步用户 LP 退出或相关 LP 状态。")
    numbered(doc, "运行 Maintenance：更新价格、全局 LP、奖励池和 burn queue。")
    numbered(doc, "运行 Claim Migrations：批量代领迁移锁仓释放，让无 UI 用户也能收到释放 KNT。")
    bullet(doc, "测试时建议先小批量执行，再执行全量，便于定位失败交易。")
    bullet(doc, "每次执行后看日志页的 tx、status 和 error 字段。")

    add_h2(doc, "3.4 Admin 常用核对入口")
    fixed_table(
        doc,
        ["场景", "Admin 核对位置", "需要看的字段"],
        [
            ["用户推荐未绑定", "财务页 / 用户详情", "referrer、推荐记录、用户是否已绑定"],
            ["USDT 入金未处理", "日志页 / USDT Deposits", "tx、confirmations、status、error"],
            ["Pancake 买卖异常", "财务页 / 交易与税费记录", "Buy/Sell 事件、税费拆分、价格"],
            ["奖励未到账", "总览页 + 日志页", "rewardPool、dailyEmission、Maintenance 记录"],
            ["迁移释放未到账", "迁移状态 + Claim Migrations 日志", "nextMigrationId、可释放数量、claim tx"],
        ],
        [1.7, 2.1, 2.6],
    )

    add_h1(doc, "4. 当前 testnet 迁移完成情况")
    fixed_table(
        doc,
        ["数据项", "结果"],
        [
            ["LP 导入", "1913 行，totalLpValueUsdt = 1360238.991159536737295168"],
            ["老算力", "65346806.123652943211440508"],
            ["节点数", "328"],
            ["推荐关系", "3558 条逐个链上校验，3558 条匹配"],
            ["迁移锁仓", "1215 条，nextMigrationId = 1216"],
            ["锁仓 KNT 总量", "351622.272577569467629156"],
            ["合约 KNT 余额", "189000000.0"],
            ["Reward Pool", "188936160.0"],
            ["最新价格", "3.0 USDT"],
            ["Global LP", "8000.0 USDT"],
        ],
        [2.0, 4.4],
    )

    add_h1(doc, "5. 常见问题")
    add_h2(doc, "推荐互转后没有绑定")
    bullet(doc, "先检查 Referral Signal KNT 是否大于 0。")
    bullet(doc, "检查两笔互转金额是否完全一致。")
    bullet(doc, "检查用户是否已经绑定过上级。")
    bullet(doc, "在 Admin 财务页输入用户地址，用链上 referrerOf 结果为准。")

    add_h2(doc, "USDT 已转入但没有入账")
    bullet(doc, "确认 USDT 合约地址是测试网 USDT：0xacD944e910952c020eb129C50921f180c62c3291。")
    bullet(doc, "确认收款地址是 KNT Proxy。")
    bullet(doc, "等待 Keeper 确认数和扫描间隔。")
    bullet(doc, "Admin 在日志页检查 Observer / Deposit 记录。")

    add_h2(doc, "用户没有 UI，如何收到释放币")
    bullet(doc, "由 Keeper/Admin 执行 Claim Migrations 批量代领。")
    bullet(doc, "合约会把可释放 KNT 直接转到用户钱包。")
    bullet(doc, "用户只需要在钱包或区块浏览器查看 KNT 到账记录。")

    add_h1(doc, "6. 操作安全要求")
    bullet(doc, "测试网和主网地址不要混用。")
    bullet(doc, "Admin 修改参数前，至少核对网络、合约地址、目标地址和权限。")
    bullet(doc, "不要把私钥、助记词、.env、Worker secret 发给任何人。")
    bullet(doc, "Keeper 钱包只保留必要权限和必要测试 BNB。")
    bullet(doc, "所有异常处理先在测试网复现，再考虑主网操作。")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("KNT Testnet Guide · 2026-05-19")
    set_run_font(run, size=9, color=TOKENS["muted"])

    doc.core_properties.title = "KNT 测试网使用教程"
    doc.core_properties.subject = "普通用户 / Admin / Keeper 快速操作手册"
    doc.core_properties.author = "KNT Project"
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
