from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "knt-test-manual.md"
OUTPUT = ROOT / "docs" / "knt-test-manual.docx"

ACCENT = "1F6F5B"
ACCENT_DARK = "145747"
MUTED = "637083"
HEADER_FILL = "E4F1ED"
TABLE_BORDER = "C6CED8"
CODE_FILL = "F3F5F7"

# The artifact renderer does not always honor inherited keep-with-next on
# Chinese heading styles, so keep known dense sections from starting at page
# bottoms with explicit breaks.
PAGE_BREAK_BEFORE_HEADINGS = {
    "4.2 BSC Testnet 业务测试",
    "9. 节点资格测试",
    "11. LP 提现和无 UI 退出测试",
}


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def set_table_width(table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths_inches: list[float]) -> None:
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def clean_inline(text: str) -> str:
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = text.replace("&gt;", ">").replace("&lt;", "<").replace("&amp;", "&")
    return text.strip()


def add_inline_runs(paragraph, text: str, font_size: float | None = None) -> None:
    parts = re.split(r"(`[^`]+`)", clean_inline(text))
    for part in parts:
        if not part:
            continue
        is_code = len(part) >= 2 and part[0] == "`" and part[-1] == "`"
        run = paragraph.add_run(part[1:-1] if is_code else part)
        set_east_asia_font(run, "Consolas" if is_code else "Microsoft YaHei")
        if font_size:
            run.font.size = Pt(font_size)
        if is_code:
            run.font.color.rgb = RGBColor(31, 47, 61)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Heading 1", 18, ACCENT_DARK, 16, 8),
        ("Heading 2", 14, ACCENT, 12, 6),
        ("Heading 3", 12, "17202C", 10, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_markdown(path: Path) -> tuple[str, list[str], list[tuple[int, str]], list]:
    lines = path.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("# ").strip()
    metadata: list[str] = []
    i = 1
    while i < len(lines) and not lines[i].startswith("## "):
        line = lines[i].strip()
        if line:
            metadata.append(line.rstrip("  "))
        i += 1

    headings: list[tuple[int, str]] = []
    blocks: list = []
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            blocks.append(("paragraph", " ".join(part.strip() for part in paragraph_lines).strip()))
            paragraph_lines = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph()
            header = split_table_row(line)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append(("table", header, rows))
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            headings.append((level, text))
            blocks.append(("heading", level, text))
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", line)
        if bullet_match:
            flush_paragraph()
            blocks.append(("bullet", bullet_match.group(1).strip()))
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if number_match:
            flush_paragraph()
            blocks.append(("number", number_match.group(1).strip(), number_match.group(2).strip()))
            i += 1
            continue

        paragraph_lines.append(line)
        i += 1

    flush_paragraph()
    return title, metadata, headings, blocks


def add_cover(document: Document, title: str, metadata: list[str]) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)

    document.add_paragraph()
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    p = document.add_paragraph()
    run = p.add_run("业务逻辑测试手册")
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("适用于合约、链上业务、keeper、后台管理台和财务统计验收。")
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    for item in metadata:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        set_paragraph_shading(p, "F7FAF9")
        add_inline_runs(p, item, 10.5)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("基线版本: BSC Testnet / KNT Proxy 0x14Dc8a0E97815128304883DEaEc89D6773937dc0")
    set_east_asia_font(run, "Consolas")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_page_break()


def add_toc(document: Document, headings: list[tuple[int, str]]) -> None:
    document.add_heading("目录", level=1)
    for level, text in headings:
        if level == 2:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            set_east_asia_font(run, "Microsoft YaHei")
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string("17202C")
        elif level == 3:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            set_east_asia_font(run, "Microsoft YaHei")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_page_break()


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.7) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    add_inline_runs(paragraph, text, font_size)
    for run in paragraph.runs:
        run.bold = bold


def table_widths(header: list[str]) -> list[float]:
    n = len(header)
    lowered = [h.lower() for h in header]
    if n == 2:
        return [2.0, 4.8]
    if n == 3 and header[0].upper() == "ID":
        return [0.65, 2.85, 3.35]
    if n == 3:
        return [1.6, 2.2, 3.0]
    if n == 4 and header[0].upper() == "ID":
        return [0.6, 2.2, 2.15, 1.9]
    if n == 4:
        return [1.25, 1.1, 2.55, 1.95]
    if n == 5:
        return [0.7, 1.35, 1.75, 1.75, 1.3]
    if n == 6:
        return [0.75, 1.25, 1.25, 1.2, 1.2, 1.2]
    return [6.8 / n] * n


def add_markdown_table(document: Document, header: list[str], rows: list[list[str]]) -> None:
    for row_index, row_data in enumerate(rows):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        if row_index % 2 == 0:
            set_paragraph_shading(p, "F7FAF9")

        for index, column in enumerate(header):
            value = row_data[index] if index < len(row_data) else ""
            if index > 0:
                p.add_run("   |   ")
            label = p.add_run(f"{clean_inline(column)}: ")
            set_east_asia_font(label, "Microsoft YaHei")
            label.bold = True
            label.font.size = Pt(9.2)
            label.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
            add_inline_runs(p, value.replace("\n", " / "), 9.2)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(document: Document, code: str) -> None:
    lines = code.splitlines() or [""]
    for index, line in enumerate(lines):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(2 if index == 0 else 0)
        p.paragraph_format.space_after = Pt(2 if index == len(lines) - 1 else 0)
        set_paragraph_shading(p, CODE_FILL)
        run = p.add_run(line)
        set_east_asia_font(run, "Consolas")
        run.font.size = Pt(8.8)
        run.font.color.rgb = RGBColor.from_string("263244")


def build_docx() -> None:
    title, metadata, headings, blocks = parse_markdown(SOURCE)
    document = Document()
    configure_styles(document)
    add_cover(document, title, metadata)
    add_toc(document, headings)

    for block in blocks:
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            heading = document.add_heading(text, level=max(1, min(level - 1, 3)))
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.keep_together = True
            if text in PAGE_BREAK_BEFORE_HEADINGS:
                heading.paragraph_format.page_break_before = True
        elif kind == "paragraph":
            _, text = block
            p = document.add_paragraph()
            add_inline_runs(p, text, 10.2)
        elif kind == "bullet":
            _, text = block
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            add_inline_runs(p, f"• {text}", 10)
        elif kind == "number":
            _, number, text = block
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            add_inline_runs(p, f"{number}. {text}", 10)
        elif kind == "code":
            _, code = block
            add_code_block(document, code)
        elif kind == "table":
            _, header, rows = block
            add_markdown_table(document, header, rows)

    for section in document.sections:
        section.top_margin = Cm(1.65)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.65)
        section.right_margin = Cm(1.65)
        footer_p = section.footer.paragraphs[0]
        footer_p.text = ""
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_inline_runs(footer_p, "KNT 业务逻辑测试手册 | 2026-05-05", 9)

    document.core_properties.title = "KNT 项目业务逻辑测试手册"
    document.core_properties.subject = "Knights-dapp 测试手册"
    document.core_properties.keywords = "KNT, Knights-dapp, 测试手册, BSC Testnet"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
